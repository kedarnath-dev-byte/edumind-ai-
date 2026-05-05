"""
@module   DOCXLoader
@description Loads and chunks Microsoft Word (.docx) files for RAG pipeline.
             Inherits BaseDocumentLoader — implements validate() and load().
             Extracts text paragraph by paragraph using python-docx library.
@author   EduMind AI Engineering
"""

import os
from typing import List

from docx import Document

from modules.ingestion.base_loader import BaseDocumentLoader, DocumentChunk


class DOCXLoader(BaseDocumentLoader):
    """
    Loads Word documents and splits them into overlapping chunks.
    Extracts text from paragraphs and tables both.
    """

    def validate(self, source: str) -> bool:
        """Check file exists and is a .docx file."""
        if not os.path.exists(source):
            raise FileNotFoundError(f"DOCX file not found: {source}")
        if not source.lower().endswith(".docx"):
            raise ValueError(f"Not a DOCX file: {source}")
        return True

    def load(self, source: str) -> List[DocumentChunk]:
        """Extract text from Word document and split into chunks."""
        self.validate(source)

        try:
            doc = Document(source)
            full_text = ""

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text += paragraph.text + "\n"

            # Extract text from tables too
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text += cell.text + "\n"

            if not full_text.strip():
                raise ValueError(f"DOCX file has no extractable text: {source}")

            chunks = self._split_into_chunks(full_text, source)
            return chunks

        except (FileNotFoundError, ValueError):
            raise
        except Exception as e:
            raise RuntimeError(f"Failed to load DOCX '{source}': {str(e)}")

    def _split_into_chunks(self, text: str, source: str) -> List[DocumentChunk]:
        """Split raw text into overlapping DocumentChunk objects."""
        chunks = []
        start = 0
        chunk_index = 0
        text_length = len(text)

        # Count total chunks first
        temp_start = 0
        total = 0
        while temp_start < text_length:
            temp_start += self.chunk_size - self.chunk_overlap
            total += 1

        # Build chunks
        while start < text_length:
            end = start + self.chunk_size
            chunk_text = text[start:end].strip()

            if chunk_text:
                chunks.append(DocumentChunk(
                    content=chunk_text,
                    source=source,
                    chunk_index=chunk_index,
                    total_chunks=total,
                    metadata={
                        "file_type": "docx",
                        "file_name": os.path.basename(source),
                        "loader": self.get_loader_name()
                    }
                ))
                chunk_index += 1

            start += self.chunk_size - self.chunk_overlap

        return chunks

